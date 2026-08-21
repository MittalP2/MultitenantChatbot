"""Build a Google Drive–ready Word report from the Lyzr RAG comparison."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "eval" / "Lyzr_Financial_RAG_Comparison_Report.docx"

NAVY = RGBColor(0x1F, 0x2A, 0x44)
GRAY = RGBColor(0x4A, 0x55, 0x68)
HEADER_FILL = "1F2A44"
ROW_ALT = "F4F6F8"
HEADER_FONT = RGBColor(0xFF, 0xFF, 0xFF)


def set_run(run, *, size=11, bold=False, color=NAVY, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "D0D5DD")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def fill_table(table, rows, header=True):
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(text)
            is_header = header and i == 0
            set_run(
                run,
                size=10,
                bold=is_header or j == 0,
                color=HEADER_FONT if is_header else NAVY,
            )
            set_cell_border(cell)
            if is_header:
                shade_cell(cell, HEADER_FILL)
            elif i % 2 == 0:
                shade_cell(cell, ROW_ALT)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, size=16 if level == 1 else 13, bold=True, color=NAVY)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run(run, size=11, color=NAVY)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run(run, size=11, color=NAVY)
    p.paragraph_format.space_after = Pt(3)
    return p


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("Financial Document Intelligence Pipeline (RAG)")
    set_run(r, size=22, bold=True, color=NAVY)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(2)
    r = sub.add_run("Chunking strategy comparison in Lyzr Studio")
    set_run(r, size=14, color=GRAY)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    r = meta.add_run("Week 2 course project  ·  Tesla, Harley-Davidson, and Polaris 10-K extracts")
    set_run(r, size=11, italic=True, color=GRAY)

    add_heading(doc, "1. Objective")
    add_body(
        doc,
        "This project builds a retrieval-augmented generation (RAG) pipeline over financial filings "
        "and compares two chunking strategies on the same questions. The experiment was run in Lyzr "
        "Studio: fixed-size chunks versus larger topical windows, with identical agents, files, and "
        "retrieval settings. The goal is to see whether chunk size changes retrieval quality and "
        "whether answers are usable for an analyst.",
    )

    add_heading(doc, "2. Corpus")
    add_body(
        doc,
        "Three SEC Form 10-K extracts were used (Item 1 Business, Item 1A Risk Factors, and Item 7 MD&A only). "
        "Full EDGAR HTML was not uploaded. Parser: PyPDF.",
    )
    add_bullet(doc, "Tesla, Inc. — TSLA_10K_2025.pdf")
    add_bullet(doc, "Harley-Davidson, Inc. — HOG_10K_2025.pdf")
    add_bullet(doc, "Polaris Inc. — PII_10K_2025.pdf")
    add_body(
        doc,
        "Files were uploaded from the project folder data/sec/lyzr_pdfs/. The same three PDFs went into both knowledge bases.",
    )

    add_heading(doc, "3. Method")
    add_body(
        doc,
        "Two Knowledge Bases and two agents were created. Role, Goal, and Instructions were the same on both agents. "
        "The only intentional difference was chunk size and overlap. Both knowledge bases were not attached to a single "
        "agent, because that would mix 800- and 1600-character windows and hide which chunker produced the hit.",
    )

    rows = [
        ["Setting", "Fixed agent", "Semantic agent"],
        ["Knowledge Base", "10k-fixed", "10k-semantic"],
        ["Chunk size / overlap", "800 / 150", "1600 / 200"],
        ["Retrieval type", "Basic", "Basic"],
        ["top-k", "5", "5"],
        ["Embedding", "text-embedding-3-small", "text-embedding-3-small"],
        ["Vector store", "Lyzr-hosted Qdrant", "Same credential"],
        ["Documents", "Same three PDFs", "Same three PDFs"],
        ["Role, Goal, Instructions", "Identical", "Identical"],
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    table.autofit = True
    fill_table(table, rows)
    doc.add_paragraph()

    add_body(
        doc,
        "Agent role: financial document analyst; answer only from the 10-K extracts. Goal: short cited answers; "
        "say you do not know if the filings do not contain the answer. Instructions: use only the knowledge base; "
        "cite the source file; do not mix companies; no investment advice.",
    )
    add_body(
        doc,
        "Reranking was not turned on. Lyzr Basic retrieval with top-k 5 is a single similarity search. This report "
        "therefore compares chunking only, not a second-stage rerank. In Lyzr, “semantic” means larger topical windows "
        "(1600 / 200), not a custom sentence-similarity splitter.",
    )
    add_body(
        doc,
        "Vector-store credentials had to be saved before training. Empty Qdrant credentials produced a 500 training "
        "error (KeyError: 'credentials'). File size was not the cause.",
    )

    add_heading(doc, "4. Evaluation")
    add_body(
        doc,
        "Six questions were asked with the same wording on the fixed agent, then on the semantic agent. "
        "Scoring used business relevance, not token overlap:",
    )
    add_bullet(doc, "Yes — an analyst could use the answer; right company; key facts present.")
    add_bullet(doc, "Partial — right company, but thin or missing a named definition.")
    add_bullet(doc, "No — wrong company, or “I don’t know” when the 10-K has the answer.")
    add_bullet(doc, "Correct refuse — “I don’t know” on a question that is not in the corpus (desired).")

    add_heading(doc, "5. Results")
    score = [
        ["#", "Question", "Fixed (800/150)", "Semantic (1600/200)", "More usable"],
        ["1", "What else does Tesla sell besides cars?", "Yes", "Yes", "Tie"],
        ["2", "Harley-Davidson dealer risks", "Yes", "Yes", "Semantic"],
        ["3", "Indian Motorcycle at Polaris", "Yes", "Yes", "Tie"],
        ["4", "What is LiveWire?", "Yes", "Yes", "Semantic"],
        ["5", "What is Autopilot?", "Partial", "Partial", "Semantic"],
        ["6", "BMW Group revenue in 2025", "Correct refuse", "Correct refuse", "Tie"],
    ]
    table = doc.add_table(rows=len(score), cols=5)
    fill_table(table, score)
    doc.add_paragraph()
    add_body(
        doc,
        "In-corpus: both agents scored 4 Yes and 1 Partial. Out-of-corpus (BMW): both correctly refused.",
    )

    add_heading(doc, "Question notes", level=2)
    add_body(
        doc,
        "1. Tesla products. Both named energy (Powerwall, Megapack, solar, Solar Roof), services, financing, "
        "in-app upgrades, and Supercharger access. Same Item 1 region.",
    )
    add_body(
        doc,
        "2. Harley dealers. Both described independent-dealer, inventory-funding, and retail-strategy risk (Item 1A). "
        "Semantic returned a clearer bullet list; fixed returned a dense paragraph of the same points.",
    )
    add_body(
        doc,
        "3. Indian Motorcycle. Both stated the 10 October 2025 agreement to sell a majority interest, close in Q1 2026, "
        "On Road reporting, held-for-sale at 31 December 2025, and impairment charges (Item 7).",
    )
    add_body(
        doc,
        "4. LiveWire. Both identified Harley’s electric brand and product types. Semantic also kept independent retail "
        "partners and a company-owned dealer, not only online D2C.",
    )
    add_body(
        doc,
        "5. Autopilot. Neither extract defined “Autopilot” by name. Both described driver-assistance, driver "
        "responsibility, and over-the-air updates. Semantic also mentioned FSD (Supervised). Neither invented a fake specification.",
    )
    add_body(
        doc,
        "6. BMW revenue. Not in the knowledge base. Fixed: “I don’t know based on the uploaded 10-K extracts.” "
        "Semantic: the same, and noted the KB holds Tesla, Harley-Davidson, and Polaris only. Neither quoted Tesla "
        "or Polaris revenue as if it were BMW.",
    )

    add_heading(doc, "6. Analysis")
    add_body(
        doc,
        "Chunk size did not change whether the right filing was found. Both agents cited the correct 10-K and item "
        "for every in-corpus question.",
    )
    add_body(
        doc,
        "Larger chunks improved how complete and readable the answer was on some questions (dealer risks, LiveWire "
        "distribution, Autopilot plus FSD). Smaller chunks were already enough for distinctive facts (Tesla product "
        "list, Indian Motorcycle sale).",
    )
    add_body(
        doc,
        "A short product question can look almost the same on both agents because embeddings land on the same paragraph. "
        "That is expected, not a failed test. This configuration also does not produce “semantic says I don’t know while "
        "fixed hallucinates.” With the same PDFs and Basic retrieval, both refused the BMW question. That is RAG behaving correctly.",
    )
    add_body(
        doc,
        "Rerank is a second ranking of a larger candidate list. It is not the same as changing top-k. It was not part of this Lyzr run.",
    )

    add_heading(doc, "7. Conclusion")
    add_body(
        doc,
        "On six identical questions, both Lyzr agents retrieved the right filings. 1600-character chunks produced slightly "
        "fuller, more structured answers on LiveWire, Autopilot, and Harley dealer risks. 800-character chunks were sufficient "
        "for Tesla products and the Indian Motorcycle sale. Out-of-corpus BMW revenue was refused by both.",
    )
    add_body(
        doc,
        "For this lexical 10-K Q&A in Lyzr, either chunk size works for fact lookup. Prefer the larger window when the user "
        "needs a whole risk section or a fuller product description.",
    )

    add_heading(doc, "8. How to reproduce")
    add_bullet(doc, "Connect Lyzr vector-store credentials (empty Qdrant credentials cause training error 500).")
    add_bullet(doc, "Create two Knowledge Bases with the settings in section 3; upload the same three PDFs to each.")
    add_bullet(doc, "Create two agents with identical Role, Goal, and Instructions; attach one KB each.")
    add_bullet(doc, "Ask the six questions on both agents and score as in section 4.")

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(18)
    r = footer.add_run("Supporting notes: eval/LYZR_COMPARISON.md  ·  Setup: README.md")
    set_run(r, size=10, italic=True, color=GRAY)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
